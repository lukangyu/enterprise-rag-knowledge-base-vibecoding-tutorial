from docx import Document
import logging
from typing import Dict, List, Any
from dataclasses import dataclass
import hashlib

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    content: str
    metadata: Dict[str, Any]
    structure: List[Dict[str, Any]]
    page_count: int
    content_hash: str


class WordParser:
    def parse(self, file_path: str) -> ParsedDocument:
        """解析Word文档"""
        doc = Document(file_path)
        content_parts = []
        structure = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)

                if para.style.name.startswith('Heading'):
                    structure.append({
                        "type": "heading",
                        "text": text,
                        "level": para.style.name
                    })

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            content_parts.append("\n".join(table_text))

        content = "\n\n".join(content_parts)
        content_hash = hashlib.sha256(content.encode()).hexdigest()

        return ParsedDocument(
            content=content,
            metadata={"file_type": "docx"},
            structure=structure,
            page_count=0,
            content_hash=content_hash
        )
